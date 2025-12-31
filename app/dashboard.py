import os
import io
import sys
import json
import tempfile
import requests
import pandas as pd
import altair as alt
from io import BytesIO
import streamlit as st
from pathlib import Path
from datetime import datetime
from docx import Document
from sqlalchemy import create_engine
from contextlib import contextmanager
from typing import Dict, List, Any
from sqlalchemy.orm import sessionmaker, Session
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from app.sql import base_sql
from app.utils import make_periods, get_filter_options, period_to_label
from app.prepare_data import hand_crafted_summary_markdown, compute_forecast_diff, table_to_nested_json, combine_projects_rows, preprocess_df_collapse_projects
from app.config import projects_list, metric_map, DB_USER, DB_PASSWORD, DB_HOST, DB_NAME, DB_DRIVER, api_key, TEMPERATURE, FREQUENCY_PENALTY, MAX_TOKENS, URL

sql = "SET NOCOUNT ON;\n" + base_sql
connection_string = (
    f"mssql+pyodbc://{DB_USER}:{DB_PASSWORD}"
    f"@{DB_HOST}/{DB_NAME}?driver={DB_DRIVER.replace(' ', '+')}"
)

engine = create_engine(connection_string)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

@contextmanager
def db_session():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# =========================================================================================
# 1) PIPELINE 
# =========================================================================================

def query_batch_to_df(db: Session, period: str) -> pd.DataFrame: 
    params = (period)
    conn = db.connection()
    raw_conn = conn.connection
    cur = raw_conn.cursor()
    cur.execute(sql, params)
    while cur.description is None:
        if not cur.nextset():
            return pd.DataFrame() 

    cols = [c[0] for c in cur.description]
    rows = cur.fetchall()
    return pd.DataFrame.from_records(rows, columns=cols)

@st.cache_data(show_spinner=False)
def run_forecast_pipeline_json(from_period, to_period, project_no, metric):
    with db_session() as db:  
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            out_paths: List[Path] = []
            try:
                for period in [from_period, to_period]:
                    df = query_batch_to_df(db, period)
                    df = combine_projects_rows(df, project_groups=projects_list, sum_cols=metric_map[metric])
                    json_file = table_to_nested_json(df, project_no)
                    out_path = tmpdir/f"output_{period}_{project_no}.json"
                    with out_path.open("w", encoding="utf-8") as f:
                        json.dump(json_file, f, ensure_ascii=False, indent=2)
                        out_paths.append(out_path)
                return compute_forecast_diff(out_paths, metric)
            except Exception as e:
                st.warning("This project data does not exist. Please check your input parameters and try again.")
                st.stop()


def chat(user_input, system_prompt):
    messages = st.session_state.messages
    prompt = system_prompt + " Question: "+ user_input
    messages.append({"role": "user", "content": user_input})
    with st.spinner("Thinking..."):
        response = requests.post(URL, data={
        "prompt": prompt,
        "temperature": TEMPERATURE,
        "frequency_penalty": FREQUENCY_PENALTY,
        "max_tokens": MAX_TOKENS}, 
        headers={"Authorization": f"Bearer {api_key}"})
        reply = response.json()['answer']
        messages.append({"role": "assistant", "content": reply})
        return reply

# =========================================================================================
# 2) HELPERS FOR DATAFRAMES
# =========================================================================================

def projects_to_dataframe(projects: Dict[str, Any]) -> pd.DataFrame:
    """One row per project, includes period1/period2 from total_<metric>."""
    rows = []
    for job_no, proj in projects.items():
        meta  = proj.get("project_meta", {}) or {}
        total = proj.get(f"total_{metric}", {}) or {}
        rows.append({
            "job_no": job_no,
            "description": meta.get("description", ""),
            "client": meta.get("client", ""),
            "period1": total.get("period1", "Previous"),
            "period2": total.get("period2", "Current"),
            "file1_metric": total.get("file1", 0.0)/1000,
            "file2_metric": total.get("file2", 0.0)/1000,
            "difference": total.get("difference", 0.0)/1000,
        })
    df = pd.DataFrame(rows)
    if not df.empty:
        df["difference_abs"] = df["difference"].abs()
    return df

def main_costtypes_df(project: Dict[str, Any]) -> pd.DataFrame:
    """Top-level items in costline_increases_trajectory."""
    traj = project.get("costline_increases_trajectory", []) or []
    df = pd.DataFrame([{
        "category": t.get("category") if not pd.isna(t.get("category")) and str(t.get("category")).strip().lower() != "nan" else "Uncategorized",
        "file1_metric": t.get(f"file1_metric", 0.0)/1000,
        "file2_metric": t.get(f"file2_metric", 0.0)/1000,
        "difference": t.get("difference", 0.0)/1000,
    } for t in traj if t.get("category")])
    return _numericize(df)


def subcategories_df(project: Dict[str, Any]) -> pd.DataFrame:
    """Flatten all subcategories under each main cost type."""
    traj = project.get("costline_increases_trajectory", []) or []
    rows = []
    for t in traj:
        parent_val = t.get("category") 
        parent = "Uncategorized" if pd.isna(parent_val) or str(parent_val).strip().lower() == "nan" else parent_val
        for s in (t.get("subcategories") or []):
            rows.append({
                "main_cost_type": parent,
                "category": "Uncategorized" if pd.isna(s.get("category")) or str(s.get("category")).strip().lower() == "nan" else s.get("category"),
                "file1_metric": s.get("file1_metric", 0.0)/1000,
                "file2_metric": s.get("file2_metric", 0.0)/1000,
                "difference": s.get("difference", 0.0)/1000,
            })
    df = pd.DataFrame(rows)
    return _numericize(df)


def children_df(project: Dict[str, Any]) -> pd.DataFrame:
    """Flatten children under each subcategory (nested by parent main + parent sub)."""
    traj = project.get("costline_increases_trajectory", []) or []
    rows = []
    for t in traj:
        parent_main = t.get("category") 
        parent_main = "Uncategorized" if pd.isna(t.get("category")) or str(parent_main).strip().lower() == "nan" else parent_main
        for s in (t.get("subcategories") or []):
            parent_sub = s.get("category") if not pd.isna(s.get("category")) and not str(s.get("category")).strip().lower() == "nan" else "Uncategorized"
            for c in (s.get("children") or []):
                rows.append({
                    "main_cost_type": parent_main,
                    "subcategory": parent_sub,
                    "category": c.get("category") if not pd.isna(c.get("category")) and not str(c.get("category")).strip().lower() == "nan" else "Uncategorized",
                    "file1_metric": c.get("file1_metric", 0.0)/1000,
                    "file2_metric": c.get("file2_metric", 0.0)/1000,
                    "difference": c.get("difference", 0.0)/1000,
                })        
    df = pd.DataFrame(rows)
    return _numericize(df)

def _numericize(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty: return df
    for col in ["file1_metric", "file2_metric", "difference"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0.0)
    return df

def build_bar_chart(df: pd.DataFrame, y_field: str, title: str) -> alt.Chart | None:
    if df.empty or y_field not in df.columns:
        return None
    chart_df = df.sort_values("difference", ascending=False)
    return (
        alt.Chart(chart_df)
        .mark_bar()
        .encode(
            x=alt.X("difference:Q", title="Difference (Current - Previous Period)"),
            y=alt.Y(f"{y_field}:N", sort="-x", title=y_field.replace("_", " ").title()),
            color=alt.Color("difference:Q", scale=alt.Scale(scheme="redyellowgreen"), title="Difference"),
            tooltip=[y_field, "file1_metric", "file2_metric", "difference"],
        )
        .properties(height=420, title=title)
        .interactive()
    )

def render_overall_summary(df_from, df_to, metric, metric_map, from_period, to_period):
    st.subheader("Overall Summary - Leaderboard")

    key_cols = ["iProjYear","cSegment","cPeriod","TYP", "cType"]
    metric_col = metric_map[metric]
    to_period, from_period = period_to_label(str(to_period)), period_to_label(str(from_period))
    col_to, col_from = f"{metric}_{to_period}", f"{metric}_{from_period}"

    @st.cache_data(show_spinner=False)
    def prep(d, metric_col, to_hash):
        d = combine_projects_rows(d, project_groups=projects_list, key_cols=key_cols, sum_cols=metric_col)
        return preprocess_df_collapse_projects(d, metric_col)

    a = prep(df_to, metric_col, to_hash=f"FROM::{to_period}::{metric_col}").rename(columns={metric_col: col_to}).set_index("iProjNo")
    b = prep(df_from, metric_col, to_hash=f"FROM::{from_period}::{metric_col}").rename(columns={metric_col: col_from}).set_index("iProjNo")[[col_from]]

    # fast join instead of merge
    out = a.join(b, how="outer").reset_index().rename(columns={
        "iProjNo": "Project", "cProjDesc": "Project Name", "cClientDesc": "Client"
    })

    # scale + diff (vectorized)
    out[[col_to, col_from]] = out[[col_to, col_from]].div(1000.0)
    out["difference"] = out[col_to].fillna(0.0) - out[col_from].fillna(0.0)

    # keep comparable rows: at least one side present AND labels present
    out = out.dropna(subset=out.columns).sort_values("difference", ascending=False, kind="mergesort").reset_index(drop=True)

    if out.empty:
        st.warning(f"No comparable projects found between {from_period.upper()} and {to_period.upper()}.")
        return

    top = out.iloc[0]
    st.info(
        f"Highest Costs Incurred from {from_period.upper()} to {to_period.upper()} "
        f"for Project {top['Project']}: {top.get('Project Name','')}") 
    numeric_cols = ["difference", col_from, col_to]
    out[numeric_cols] = out[numeric_cols].round(3)
    st.dataframe(out, use_container_width=True)
    
def render_trajectory_tree(project: Dict[str, Any], period1_label: str, period2_label: str):
    """
    Expandable tree UI:
      Main Cost Type → Subcategory → Sub-Subcategory
    """
    traj = project.get("costline_increases_trajectory", []) or []
    if not traj:
        st.info("No trajectory data for this project.")
        return

    for main in traj:
        main_cat = main.get("category") if not pd.isna(main.get("category")) else "Uncategorized"
        m1 = main.get("file1_metric", 0.0) or 0.0
        m2 = main.get("file2_metric", 0.0) or 0.0
        md = main.get("difference", 0.0) or 0.0

        with st.expander(f"▶ {main_cat}  |  Δ {md/1000:,.2f}", expanded=False):
            c1, c2, c3 = st.columns(3)
            c1.metric(period1_label, f"{m1/1000:,.2f}")
            c2.metric(period2_label, f"{m2/1000:,.2f}")
            c3.metric("Difference", f"{md/1000:,.2f}")

            subs = main.get("subcategories") or []
            if not subs:
                st.caption("No subcategories.")
                continue

            for sub in subs:
                sub_cat = sub.get("category") or "Uncategorized"
                s1 = sub.get("file1_metric", 0.0) or 0.0
                s2 = sub.get("file2_metric", 0.0) or 0.0
                sd = sub.get("difference", 0.0) or 0.0

                with st.expander(f"↳ {sub_cat}  |  Δ {sd/1000:,.2f}", expanded=False):
                    sc1, sc2, sc3 = st.columns(3)
                    sc1.metric(period1_label, f"{s1/1000:,.2f}")
                    sc2.metric(period2_label, f"{s2/1000:,.2f}")
                    sc3.metric("Difference", f"{sd/1000:,.2f}")

                    kids = sub.get("children") or []
                    if not kids:
                        st.caption("No children.")
                        continue

                    # Children table inside this expander
                    kids_df = pd.DataFrame([{
                        "category": k.get("category"),
                        period1_label: k.get("file1_metric", 0.0)/1000,
                        period2_label: k.get("file2_metric", 0.0)/1000,
                        "difference": k.get("difference", 0.0)/1000,
                    } for k in kids if k.get("category")])

                    if not kids_df.empty:
                        kids_df = _numericize(kids_df)
                        kids_df = kids_df.sort_values("difference", ascending=False)
                        st.dataframe(kids_df) #, use_container_width=True, height=240
    
# ===========================================================================================
# 3) STREAMLIT UI
# ==========================================================================================

st.set_page_config(layout="wide")
st.markdown(
    """
    <div style="
        background-color:#0f172a;
        padding:20px 20px 5px 20px;
        border-radius:8px;
        border:1px solid #1e293b;
        margin-bottom:15px;">
        <h1 style="color:#4DB8FF; font-weight:800; margin:0;">
            Finance: Cost Monitoring Dashboard 
        </h1>
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """ 
This dashboard visualizes the differences in project costs (EAC or YTD), highlighting project-level changes, major cost-type contributions, 
sub cost lines changes under each main cost type and further down to the lowest level of cost line under each sub cost line. Values are in millions. 
"""
)
st.sidebar.image(r"C:\Users\amrin.kareem\Downloads\Finance\assets\nmdc_logo.png", width="content")


# -------- sidebar UI --------
st.sidebar.header("Choose Project Details")

# Period dropdowns (YYYYMM)
with db_session() as db:  
    periods, projects = get_filter_options(db)
    
#check if format of dates is %Y%m
periods = [i for i in periods if isinstance(i, str) and len(i) == 6 and i.isdigit()]
period_options = make_periods(start_ym=min(periods), end_ym=max(periods))  # set end_ym to today
period_labels = {p: period_to_label(p) for p in period_options}

from_period = st.sidebar.selectbox(
    "From period",
    options=period_options,
    index=period_options.index("201901") if "201901" in period_options else datetime.now().strftime("%Y%m"),
    format_func=lambda p: f"{period_labels[p]}",
)

to_period = st.sidebar.selectbox(
    "To period",
    options=period_options,
    index=period_options.index("202001") if "202001" in period_options else datetime.now().strftime("%Y%m"),
    format_func=lambda p: f"{period_labels[p]}",
)
if int(to_period) < int(from_period):
    st.sidebar.error("To period must be the same as or after From period.")

project_no = st.sidebar.selectbox(
    "Project number",
    options=projects,
    index=projects.index(2208),
)

metric = st.sidebar.selectbox(
    "Metric",
    options=["forecast_costs_at_completion", "ytd_actual"],
    index=0,
)

run_button = st.sidebar.button("Run analysis", type="primary")
if "result" not in st.session_state:
    st.session_state.result = None
if "analysis_done" not in st.session_state:
    st.session_state.analysis_done = False
if "project_chats" not in st.session_state:
    st.session_state.project_chats = {}

if run_button:
    if not from_period or not to_period:
        st.sidebar.error("Please choose the period before running.")
    with st.spinner("Running forecast comparison..."):
        try:
            
            st.session_state.result = run_forecast_pipeline_json(
                from_period, to_period,
                project_no, 
                metric
            )
            st.session_state.analysis_done = True
            
        except Exception as e:
            st.session_state.result = None
            st.session_state.analysis_done = False
            st.exception(e)

if not st.session_state.analysis_done or st.session_state.result is None:
    st.info("Enter the to and from period and run analysis to get results.")
    st.stop()

result = st.session_state.result
projects = result.get("projects", {})
try:
    parts = hand_crafted_summary_markdown(projects, metric)
except Exception as e:
    st.warning(f"Data collected for '{metric}'. Click **Run Analysis** to compute.")
    

projects_df = projects_to_dataframe(projects)
# Period labels for slope chart
period1_label = projects_df["period1"].iloc[0]
period2_label = projects_df["period2"].iloc[0]

overall_summary, tab_summary, tab_main, tab_sub, tab_children, tab_drilldown = st.tabs(["Overall Summary", "Project Summary", "Main Cost Driver", "Sub Categories", "Sub-Subcategories", "Cost Trajectory Explorer"])
# =========================================================================================
# TAB 0: OVERALL SUMMARY
# =========================================================================================
with db_session() as db:
        df_from = query_batch_to_df(db, from_period)
        df_to = query_batch_to_df(db, to_period)
with overall_summary:
    render_overall_summary(df_from, df_to, metric, metric_map, from_period, to_period)

# =========================================================================================
# TAB 1: SUMMARY
# =========================================================================================
with tab_summary:
    st.subheader("Project Summary")
    selected_job = st.selectbox("Selected Project", projects_df["job_no"].tolist())
    project = projects[selected_job]

    meta = project.get("project_meta", {}) or {}
    total = project.get(f"total_{metric}", {}) or {}
    st.write(f"**{selected_job} — {meta.get('description','')} ({meta.get('client','')})**")
    st.metric("Project difference", f"{total.get('difference',0.0)/1000:,.2f} million")
    
    ##############  Summary Generation ################
    doc = Document()
    doc.add_heading('Cost Summary', level=1)
    doc.add_paragraph(parts)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create a download button for the Word document
    st.download_button(
        label="Download Summary",
        data=buffer,
        file_name="Analysis.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    c2, c3 = st.columns(2)
    c2.metric(f"Total {period1_label} {metric}", f"{projects_df['file1_metric'].sum():,.2f} million")
    c3.metric(f"Total {period2_label} {metric}", f"{projects_df['file2_metric'].sum():,.2f} million",
        delta=f"{projects_df['difference'].sum():,.2f} million"
    )
    # Slope chart per project
    long_df = projects_df.melt(
        id_vars=["job_no", "description", "client"],
        value_vars=["file1_metric", "file2_metric"],
        var_name="period_key",
        value_name="metric_value"
    )
    
    long_df["period"] = long_df["period_key"].replace({
        "file1_metric": period1_label,
        "file2_metric": period2_label
    })

    slope_chart = (
        alt.Chart(long_df)
        .mark_line(point=True, strokeWidth=3)
        .encode(
            x=alt.X("period:N", title="Period", sort=[period1_label, period2_label]),
            y=alt.Y("metric_value:Q", title=f"{metric}"),
            color=alt.Color("job_no:N", legend=None),
            tooltip=["job_no", "description", "client", "period", "metric_value"]
        )
        .properties(height=450)
        .interactive()
    )
    st.altair_chart(slope_chart, use_container_width=True)

########### Download Cost Breakdown Report as Excel ###########

    job_no_first = projects_df["job_no"].tolist()[0]
    df_projects_sorted = projects_df.sort_values(by="difference_abs", ascending=False)

    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_projects_sorted.to_excel(writer, sheet_name="Projects", index=False)

        for job_no, proj in projects.items():
            df_costlines = subcategories_df(proj)
            df_main_types = main_costtypes_df(proj)
            df_children = children_df(proj)

            if not df_main_types.empty:
                df_main_types.sort_values("difference", ascending=False)\
                    .to_excel(writer, sheet_name=f"{job_no}_main_types"[:31], index=False)

            if not df_costlines.empty:
                df_costlines.sort_values("difference", ascending=False)\
                    .to_excel(writer, sheet_name=f"{job_no}_costlines"[:31], index=False)

            if not df_children.empty:
                df_children.sort_values("difference", ascending=False)\
                    .to_excel(writer, sheet_name=f"{job_no}_children"[:31], index=False)

    data = buffer.getvalue()

    st.download_button(
        label="Download Cost Breakdown Report",
        data=data,
        file_name=f"{metric}_diff_{job_no_first}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    ############ Project-wise Chat ################
    st.subheader("Ask Questions about this Project")
    if project:
        total = project.get(f"total_{metric}", {}) or {}
        period1_label = total.get("period1", "previous period")
        period2_label = total.get("period2", "current period")

        project_key = f"{selected_job}|{period1_label}|{period2_label}"

        # system_prompt = SYSTEM_PROMPT_TEMPLATE.format(
        #     period1_label=period1_label,
        #     period2_label=period2_label,
        #     parts=parts)
        # system_prompt = "Read the following financial expense report carefully and answer the questions." + parts
        system_prompt = (
            "You are a financial analyst. Use ONLY the provided report context. "
            "Answer with numbers and name the costline/category.\n\n"
            "REPORT CONTEXT:\n" + parts
        )

        if project_key not in st.session_state.project_chats:
            st.session_state.project_chats[project_key] = [{"role": "system", "content": system_prompt}]
        st.session_state.messages = st.session_state.project_chats[project_key]
        if st.button("Which is the main cost contributor for this project?"):
            reply = chat("Which is the main cost contributor for this project?", system_prompt)
            st.session_state.project_chats[project_key] = st.session_state.messages
        prompt = st.chat_input(f"Which is the main cost contributor for this project?") # Ask about project {selected_job} ({period1_label} → {period2_label})
        if prompt:
            reply = chat(prompt, system_prompt)
            st.session_state.project_chats[project_key] = st.session_state.messages

        for message in st.session_state.project_chats[project_key][1:]:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if st.button("Reset Conversation"):
            st.session_state.project_chats[project_key] = [{"role": "system", "content": system_prompt}]
            st.success(f"Conversation for project {selected_job} reset.")
    else:
        st.info("Select a project to start chatting about it.")
            
# =========================================================================================
# TAB 2: MAIN COST TYPES
# =========================================================================================
with tab_main:
    st.subheader("Main Cost Type Breakdown")
    selected_job = st.selectbox("Selected Project", projects_df["job_no"].tolist(), key="main_job")
    project = projects[selected_job]

    main_df = main_costtypes_df(project)
    if main_df.empty:
        st.info("No main cost types found.")
    else:
        chart = build_bar_chart(main_df, "category", "Main Cost Type Differences")
        if chart: st.altair_chart(chart, use_container_width=True)
        st.dataframe(main_df, use_container_width=True)

# =========================================================================================
# TAB 3: SUBCATEGORY BREAKDOWN
# =========================================================================================
with tab_sub:
    st.subheader("Cost Subcategory Breakdown")
    selected_job_sub = st.selectbox("Selected Project", projects_df["job_no"].tolist(), key="sub_job")
    project_sub = projects[selected_job_sub]
    sub_df = subcategories_df(project_sub)
    if sub_df.empty:
        st.info("No subcategories found.")
    else:
        parent_options = ["ALL"] + sorted(sub_df["main_cost_type"].dropna().unique().tolist())
        parent_choice = st.selectbox("Filter by Main Cost Type", parent_options)
        view_df = sub_df if parent_choice == "ALL" else sub_df[sub_df["main_cost_type"] == parent_choice]
        chart = build_bar_chart(view_df, "category", f"Subcategory Differences ({parent_choice})")
        if chart: st.altair_chart(chart, use_container_width=True)
        st.dataframe(view_df, use_container_width=True)

# =========================================================================================
# TAB 4: CHILDREN BREAKDOWN
# =========================================================================================
with tab_children:
    st.subheader("Sub-Subcategory-Level Breakdown")

    selected_job_ch = st.selectbox("Selected Project", projects_df["job_no"].tolist(), key="child_job")
    project_ch = projects[selected_job_ch]
    ch_df = children_df(project_ch)
    if ch_df.empty:
        st.info("No children found.")
    else:
        main_opts = ["ALL"] + sorted(ch_df["main_cost_type"].dropna().unique().tolist())
        main_choice = st.selectbox("Filter by Main Cost Type", main_opts, key="child_main")

        filtered = ch_df if main_choice == "ALL" else ch_df[ch_df["main_cost_type"] == main_choice]

        sub_opts = ["ALL"] + sorted(filtered["subcategory"].dropna().unique().tolist())
        sub_choice = st.selectbox("Filter by Subcategory", sub_opts, key="child_sub")

        view_df = filtered if sub_choice == "ALL" else filtered[filtered["subcategory"] == sub_choice]
        mask = ch_df["category"].astype(str).str.contains("Income Tax", na=False)
        for _, c in ch_df[mask].iterrows():
            st.info(f"""**INCOME TAX DATA**:
                    
                {period1_label.upper()}: {c.get('file1_metric', 0.0):.2f} million
                {period2_label.upper()}: {c.get('file2_metric', 0.0):.2f} million   
                Difference: {c.get('difference', 0.0):.2f} million
                """)
        chart = build_bar_chart(view_df, "category", f"Children Differences ({main_choice} → {sub_choice})")
        if chart: st.altair_chart(chart, use_container_width=True)

        st.dataframe(view_df, use_container_width=True)
        

# =========================================================================================
# TAB 5: DRILLDOWN
# =========================================================================================
with tab_drilldown:
    st.subheader("Expandable Cost Tree")

    selected_job_tree = st.selectbox(
        "Selected Project",
        projects_df["job_no"].tolist(),
        key="tree_job"
    )
    project_tree = projects[selected_job_tree]

    meta = project_tree.get("project_meta", {}) or {}
    st.write(f"**{selected_job_tree} — {meta.get('description','')} ({meta.get('client','')})**")

    render_trajectory_tree(project_tree, period1_label, period2_label)

